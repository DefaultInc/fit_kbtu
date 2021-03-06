from datetime import datetime

from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse, HttpResponseRedirect
from django.shortcuts import render, get_object_or_404
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from doc_app.forms import ClaimForm
from doc_app.models import Claim


# Create your views here.
@login_required(login_url='/auth/login/')
def document_print_view(request, id=id):
    claim = get_object_or_404(Claim, id=id)
    document = Document()
    document.add_paragraph('Декану ФИТ').alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    document.add_paragraph('Акжаловой А. Ж.').alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    document.add_paragraph(str.format('ФИТ, {} {} курса', request.user.speciality, request.user.study_year)) \
        .alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    document \
        .add_paragraph(str.format("{} {} \n\n\n", request.user.first_name, request.user.last_name)) \
        .alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    # document.add_paragraph('')
    # document.add_paragraph('')
    # document.add_paragraph('')
    document.add_heading('Заявление', level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.add_paragraph('').alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    document.add_paragraph(str.format('\t\t\t{}', claim.content))
    document.add_paragraph(str.format('\n \t\t\tПодпись \t\t\t\t Дата \n \t\t\t\t\t\t\t\t {}.{}.{}',
                                      datetime.now().day, datetime.now().month, datetime.now().year))\
                           .alighment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    document.save(response)
    response['Content-Disposition'] = 'attachment; filename=document.docx'
    return response


@login_required(login_url='/auth/login/')
def document_create_view(request):
    form = ClaimForm(request.POST or None, request.FILES or None)
    if form.is_valid():
        instance = form.save(commit=False)
        instance.save()
        # message success
        messages.success(request, "Successfully Created!")
        return HttpResponseRedirect(instance.get_absolute_url())
    context_dict = {
        "form": form,
    }
    context = context_dict.items()
    return render(request, "doc_app/document_create.html", context)


@login_required(login_url='/auth/login/')
def document_sign_view(request):
    return


def document_list_view(request):
    claims = Claim.objects.all()
    context = {
        "claims": claims,
    }
    return render(request, "doc_app/document_list.html", context)


@login_required(login_url='/auth/login/')
def document_detail_view(request, id=None):
    """
    post details function,
    taking id or 404 if it is not exist, give the title and instance to the view.
    """
    instance = get_object_or_404(Claim, id=id)
    context = {
        "title": instance.title,
        "instance": instance,
    }
    return render(request, "doc_app/document_detail.html", context)
